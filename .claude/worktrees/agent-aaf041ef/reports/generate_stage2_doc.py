from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DESKTOP = "/Users/annita/Desktop/stage2_deep_dive.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

NAVY   = RGBColor(0x1F, 0x38, 0x64)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
TEAL   = RGBColor(0x17, 0xB8, 0x97)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x26, 0x26, 0x26)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = TEAL

def para(text, bold_parts=None):
    p = doc.add_paragraph()
    if bold_parts is None:
        p.add_run(text)
    else:
        for segment, is_bold in bold_parts:
            r = p.add_run(segment)
            r.bold = is_bold
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent = Inches(0.4)

def add_table(headers, rows, col_widths=None, header_color="1F3864"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        shade_cell(hdr[i], header_color)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = val
            shade_cell(cells[ci], bg)
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════
h1("Stage 2 Deep Dive — Meta-Learning for Contract Renegotiation")
para("Signal Fusion with Meta-Learners | Master Thesis\nThis document explains the full Stage 2 pipeline: "
     "structure, episodic task construction, support/query mechanics, each method, and all output files.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 1. BIG PICTURE
# ══════════════════════════════════════════════════════════
h2("1. Big Picture — What Stage 2 Is Trying to Do")
para("Stage 1 trains a general-purpose MLP on all contracts using weak (Snorkel) supervision. "
     "That model knows about contracts in general but has never seen Logistics-specific patterns.")
doc.add_paragraph()
para("Stage 2 asks: given only a handful of labelled Logistics contracts (the support set), "
     "can we adapt the Stage 1 MLP so it generalises well to unseen Logistics contracts (the query set)?")
doc.add_paragraph()
para("This is a few-shot learning problem. The answer is meta-learning: instead of training a model "
     "from scratch each time, we train the Stage 1 weights to be a good starting point — weights that "
     "can be quickly adapted to any new department with very few labelled examples.")

# ══════════════════════════════════════════════════════════
# 2. TASKS = DEPARTMENTS
# ══════════════════════════════════════════════════════════
h2("2. Tasks = Departments")
para("In standard few-shot meta-learning, each 'task' is a classification problem sampled from a "
     "distribution of tasks. In this thesis:")
doc.add_paragraph()

add_table(
    ["Concept", "In this thesis"],
    [
        ("Task distribution",    "The set of all procurement departments (14 total)"),
        ("One task",             "One department (e.g. Packaging Material, Bioprocessing, ...)"),
        ("Meta-train tasks",     "All departments EXCEPT Logistics (~9 valid departments)"),
        ("Meta-test task",       "Logistics — the target department, held out entirely"),
        ("Label used",           "gold_y (binary: 0 = no renegotiation, 1 = renegotiation)"),
        ("Features",             "~160 engineered columns from contracts_with_features_labeled_with_gold.csv"),
    ],
    col_widths=[2.2, 4.8],
)

para("A department is considered a valid meta-train task only if it has enough labelled contracts "
     "to form both a support set and a non-empty query set. Departments that fail this check are "
     "excluded. This is handled by filter_valid_departments() in episode_sampler.py.")

# ══════════════════════════════════════════════════════════
# 3. EPISODES, SUPPORT, QUERY
# ══════════════════════════════════════════════════════════
h2("3. Episodes, Support Sets, and Query Sets")

h3("3.1 What is an Episode?")
para("An episode is one instantiation of a few-shot task. For each episode you:")
bullet("Pick a department")
bullet("Sample n_support_pos positive contracts + n_support_neg negative contracts → Support Set")
bullet("The remaining contracts form the → Query Set")
bullet("Adapt the model on the support set (inner loop)")
bullet("Evaluate on the query set (outer loop loss or evaluation metric)")
doc.add_paragraph()
para("Crucially, splitting is done at the contract level — the same contract_id never appears "
     "in both support and query. Since each contract can have multiple rows (one per observation year), "
     "all rows of a chosen support contract go into the support set together.")

h3("3.2 Support Set")
add_table(
    ["Property", "Detail"],
    [
        ("Size",       "n_support_pos positive contracts + n_support_neg negative contracts"),
        ("Default k",  "k=2 (2 pos + 2 neg), also tested with k=5 and k=10"),
        ("Purpose",    "The few labelled examples the model uses to adapt to this department"),
        ("Used for",   "Inner loop gradient steps (adaptation)"),
        ("Key point",  "Sampled at contract level — all years of a contract stay together"),
    ],
    col_widths=[2.0, 5.0],
)

h3("3.3 Query Set")
add_table(
    ["Property", "Detail"],
    [
        ("Size",       "All remaining contracts after support selection"),
        ("Purpose",    "Unseen contracts — what we actually want to predict"),
        ("Used for",   "Computing outer loop loss (meta-train) or evaluation metrics (meta-test)"),
        ("Key point",  "No overlap with support at the contract_id level"),
        ("Evaluation", "AUROC, F1, Average Precision, NDCG@5/10/20 computed here"),
    ],
    col_widths=[2.0, 5.0],
)

h3("3.4 Meta-Training vs Meta-Test Episodes")
add_table(
    ["", "Meta-Train Episodes", "Meta-Test Episodes (Logistics)"],
    [
        ("Departments",   "~9 valid source departments",        "Logistics only"),
        ("Repeats",       "200 iterations × batch of 4",        "20 repeated episodes (different random seeds)"),
        ("Purpose",       "Update the shared MLP weights",      "Final evaluation — report these metrics"),
        ("Loss computed", "Query loss → backprop outer loop",   "AUROC, F1, AP, NDCG@k"),
        ("Logistics?",    "Never seen during meta-training",    "Only seen at evaluation time"),
    ],
    col_widths=[1.8, 2.6, 2.6],
)
para("The 20 repeated Logistics episodes use different random seeds (42, 43, ..., 61) so that "
     "reported metrics are averages over different support/query splits — this gives more stable estimates.")

# ══════════════════════════════════════════════════════════
# 4. STAGE 1 INITIALIZATION
# ══════════════════════════════════════════════════════════
h2("4. Stage 1 Initialization")
para("Every Stage 2 experiment starts from a pretrained Stage 1 MLP, not random weights. "
     "This is the key design choice: Stage 1 has already learned contract risk patterns from "
     "all 9,201 contract-year rows using weak supervision. Stage 2 fine-tunes this knowledge "
     "toward fast department adaptation.")
doc.add_paragraph()
add_table(
    ["Init name", "Stage 1 training signal", "Model path"],
    [
        ("A_weak_only", "Snorkel soft labels only (renegotiation_prob)",           "models/stage_1/A_weak_only/mlp_pretrained.pt"),
        ("C_hybrid",    "Snorkel + gold labels with sample weighting",             "models/stage_1/C_hybrid/mlp_pretrained.pt"),
    ],
    col_widths=[1.5, 3.2, 3.3],
)
para("The preprocessor (mlp_pretrained_preprocessor.joblib) is always loaded alongside the weights — "
     "it handles imputation, scaling, and encoding consistently between Stage 1 and Stage 2.")

# ══════════════════════════════════════════════════════════
# 5. THE FOUR METHODS
# ══════════════════════════════════════════════════════════
h2("5. The Four Methods")

h3("5.1 finetune (Fine-Tuning Baseline)")
para("Not a meta-learning method. Simply takes the Stage 1 MLP and fine-tunes all its weights "
     "on the Logistics support set using gradient descent. No inner/outer loop distinction.")
bullet("Inner loop: gradient steps on ALL parameters using support set loss")
bullet("No meta-training on source departments")
bullet("Serves as the simplest possible adaptation baseline")
bullet("Source: stage2.py → run_finetune_baseline()")

h3("5.2 ANIL — Almost No Inner Loop")
para("Adapts ONLY the final classification head during the inner loop. The feature extractor "
     "(all layers except the last linear layer) is frozen during adaptation.")
bullet("Meta-train outer loop: updates ALL parameters (including feature extractor) via Adam")
bullet("Inner loop at adapt time: gradient steps on head_weight and head_bias only")
bullet("Features are extracted once with no_grad, then only the linear head is updated")
bullet("Fastest adaptation — only 2 parameters updated per inner step")
bullet("Source: src/master_thesis/anil.py → adapt_anil_on_episode(), meta_train_anil()")
doc.add_paragraph()
code_block("head_weight = head_weight - inner_lr * grad_head_weight\nhead_bias   = head_bias   - inner_lr * grad_head_bias")

h3("5.3 FOMAML — First-Order MAML")
para("Adapts ALL model parameters in the inner loop, but drops second-order gradient terms "
     "for computational efficiency. This is Reptile-adjacent — the outer gradient is simply "
     "the gradient at the adapted parameters, not the gradient-of-gradient.")
bullet("Inner loop: gradient steps on ALL parameters using support loss (create_graph=False)")
bullet("Outer loop: backprop through the adapted params WITHOUT second-order terms")
bullet("Much cheaper than full MAML — no Hessian computation")
bullet("Source: src/master_thesis/fomaml.py → adapt_fomaml_on_episode(), meta_train_fomaml()")
doc.add_paragraph()
code_block("grads = torch.autograd.grad(loss, params.values(), create_graph=False)\n# outer step uses these first-order gradients directly")

h3("5.4 MAML — Full Second-Order MAML")
para("The complete MAML algorithm. Adapts ALL parameters in the inner loop and backpropagates "
     "through the inner loop gradient computation, producing true second-order gradients.")
bullet("Inner loop: gradient steps on ALL parameters using support loss (create_graph=True)")
bullet("Outer loop: loss computed on query set with adapted params, then backward() through inner steps")
bullet("Uses torch.func.functional_call() for clean functional parameter updates")
bullet("Most expressive but most expensive — requires storing the full computation graph through inner steps")
bullet("Source: src/master_thesis/maml.py → adapt_maml_on_episode(), meta_train_maml()")
doc.add_paragraph()
code_block("grads = torch.autograd.grad(loss, params.values(), create_graph=True)  # 2nd order\nparams = {name: param - inner_lr * grad ...}")

h3("Method Comparison Table")
add_table(
    ["Method", "What adapts in inner loop", "2nd order grads", "Meta-trained?", "Speed"],
    [
        ("finetune", "All params",       "No",  "No",  "Fastest"),
        ("ANIL",     "Head only",        "Yes", "Yes", "Fast"),
        ("FOMAML",   "All params",       "No",  "Yes", "Medium"),
        ("MAML",     "All params",       "Yes", "Yes", "Slowest"),
    ],
    col_widths=[1.4, 2.3, 1.5, 1.4, 1.4],
)

# ══════════════════════════════════════════════════════════
# 6. THE TRAINING LOOP
# ══════════════════════════════════════════════════════════
h2("6. The Meta-Training Loop (Step by Step)")
para("For ANIL, FOMAML, and MAML the outer loop runs for meta_iterations=200 iterations:")
doc.add_paragraph()

steps = [
    ("Step 1 — Sample a meta-batch",
     "Randomly sample meta_batch_size=4 departments from the valid source departments (not Logistics)."),
    ("Step 2 — For each department, sample an episode",
     "Call sample_support_query_split(): pick n_support_pos + n_support_neg contracts for support, "
     "rest go to query. Contract-level — no leakage."),
    ("Step 3 — Inner loop (task adaptation)",
     "Run inner_steps gradient updates on the support set loss. What gets updated depends on the method "
     "(head only for ANIL, all params for FOMAML/MAML)."),
    ("Step 4 — Compute query loss with adapted params",
     "Evaluate the adapted model on the query set. This is the signal for the outer loop."),
    ("Step 5 — Outer loop update",
     "Average query losses across the batch. Call backward() and update the shared model via Adam "
     "(outer_lr=0.001). Gradient clipping at norm=1.0."),
    ("Step 6 — Log history",
     "Record iteration, meta_loss, n_skipped, n_valid_tasks → saved to history.csv."),
]
for title, detail in steps:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(title + ": ")
    r.bold = True
    p.add_run(detail)

# ══════════════════════════════════════════════════════════
# 7. EVALUATION ON LOGISTICS
# ══════════════════════════════════════════════════════════
h2("7. Evaluation on Logistics (Meta-Test)")
para("After meta-training, the model is evaluated on Logistics using 20 repeated episodes "
     "(make_logistics_meta_test_split(), seeds 42–61). For each episode:")
bullet("Support set: n_support_pos + n_support_neg Logistics contracts (labelled)")
bullet("Query set: remaining Logistics contracts (unseen)")
bullet("Inner loop adaptation runs target_inner_steps steps on the support set")
bullet("Predictions are made on the query set → AUROC, F1, AP, NDCG@5/10/20 computed")
bullet("Metrics are averaged across all 20 episodes → mean ± std reported")
doc.add_paragraph()
para("This repeated evaluation strategy is important: with so few Logistics contracts, a single "
     "support/query split would be highly sensitive to which specific contracts were picked. "
     "20 repeats gives a stable estimate of adaptation performance.")

# ══════════════════════════════════════════════════════════
# 8. EXPERIMENT GRID
# ══════════════════════════════════════════════════════════
h2("8. Experiment Grid")
para("Stage 2 sweeps over a grid of hyperparameters. Each combination is one experiment, "
     "saved to its own folder with a descriptive ID.")
doc.add_paragraph()
add_table(
    ["Hyperparameter", "Values tested", "What it controls"],
    [
        ("method",         "finetune, anil, fomaml, maml",  "Which adaptation algorithm"),
        ("n_support_pos",  "2, 5, 10",                      "Positive contracts in support set"),
        ("n_support_neg",  "2, 5, 10",                      "Negative contracts in support set"),
        ("inner_steps",    "1, 3, 5, 10",                   "Gradient steps in inner loop"),
        ("init_name",      "A_weak_only, C_hybrid",         "Which Stage 1 MLP to start from"),
        ("target_dept",    "Logistics",                     "Always Logistics (held-out target)"),
    ],
    col_widths=[1.8, 2.2, 3.0],
)
para("Experiment ID format:")
code_block("stage2__init-A_weak_only__method-anil__kpos-2__kneg-2__steps-3__target-Logistics")

# ══════════════════════════════════════════════════════════
# 9. OUTPUT FILES
# ══════════════════════════════════════════════════════════
h2("9. Output Files — What Each File Contains")
para("Every experiment is saved under:")
code_block("models/stage_2/experiments/{experiment_id}/")
doc.add_paragraph()

add_table(
    ["File", "Contents", "Used by"],
    [
        ("resolved_config.yaml",
         "Exact configuration snapshot used for this experiment — all hyperparameters, paths, grid values",
         "Reproducibility, debugging"),
        ("{method}_initialized.pt",
         "PyTorch weights of the meta-trained model after all outer loop iterations",
         "Resuming, further analysis"),
        ("metrics.csv",
         "Per-episode evaluation metrics: episode_idx, AUROC, F1, Average Precision, NDCG@5/10/20, log-loss — one row per Logistics episode",
         "17.1 dashboard, thesis tables"),
        ("predictions.csv",
         "Per-contract predictions across all episodes: episode_idx, repeat_idx, contract_id, department, y_true, y_prob, method",
         "Calibration analysis, error analysis"),
        ("history.csv",
         "Meta-training history: iteration, meta_loss (outer loop), n_skipped, n_valid_tasks — one row per logged iteration",
         "Adaptation curve plots"),
        ("department_summary.csv",
         "Gold label counts per department: n_pos_contracts, n_neg_contracts, pos_rate — shows task diversity",
         "Diagnostic, validity check"),
        ("department_validity.csv",
         "Per-department validity flags: valid_department (True/False), validity_reason — explains which depts were excluded and why",
         "Diagnostic, debugging"),
        ("stage2_result_summary.json",
         "Aggregated result for this experiment: mean ± std of all metrics across 20 Logistics episodes",
         "experiment_summary.csv registry"),
    ],
    col_widths=[2.2, 3.6, 1.8],
)

h3("Master Registry")
add_table(
    ["File", "Contents"],
    [
        ("models/stage_2/experiments/experiment_summary.csv",
         "One row per experiment. Columns: experiment_id, method, init_name, n_support_pos, n_support_neg, "
         "inner_steps, target_department, status, gold_auroc_mean, gold_auroc_std, gold_f1_mean, "
         "gold_ap_mean, ndcg_at_10_mean, ... — the single file read by notebook 17.1"),
    ],
    col_widths=[3.5, 4.5],
)

# ══════════════════════════════════════════════════════════
# 10. ORCHESTRATION CHAIN
# ══════════════════════════════════════════════════════════
h2("10. Orchestration Chain")
code_block(
    "stage2_config.yaml\n"
    "   ↓\n"
    "Notebook 17  OR  experiments/run_stage2.py (CLI)\n"
    "   ↓\n"
    "src/master_thesis/stage2.py   ← coordinator: resolves grid, loads data, dispatches method\n"
    "   ├── episode_sampler.py     ← builds support/query episodes per department\n"
    "   ├── anil.py                ← ANIL meta-train + evaluate\n"
    "   ├── fomaml.py              ← FOMAML meta-train + evaluate\n"
    "   ├── maml.py                ← full MAML meta-train + evaluate\n"
    "   └── metrics.py             ← AUROC, AP, F1, NDCG@k\n"
    "   ↓\n"
    "models/stage_2/experiments/{id}/   ← all artifacts saved here\n"
    "   ↓\n"
    "models/stage_2/experiments/experiment_summary.csv   ← master registry updated\n"
    "   ↓\n"
    "Notebook 17.1   ← reads summary only, generates thesis tables"
)

# ══════════════════════════════════════════════════════════
# 11. KEY PARAMETERS
# ══════════════════════════════════════════════════════════
h2("11. Key Hyperparameters at a Glance")
add_table(
    ["Parameter", "Default value", "Where set"],
    [
        ("meta_iterations",     "200",       "stage2_config.yaml → base_config.meta"),
        ("meta_batch_size",     "4",         "stage2_config.yaml → base_config.meta"),
        ("inner_lr",            "0.01",      "stage2_config.yaml → base_config.meta"),
        ("outer_lr",            "0.001",     "stage2_config.yaml → base_config.meta"),
        ("inner_steps",         "1/3/5/10",  "stage2_config.yaml → experiment_grid"),
        ("n_support_pos/neg",   "2/5/10",    "stage2_config.yaml → experiment_grid"),
        ("n_repeats (Logistics)","20",       "stage2.py → make_logistics_meta_test_split()"),
        ("target_inner_steps",  "5",         "stage2_config.yaml → base_config.task"),
        ("MLP hidden dims",     "256 → 128", "stage2_config.yaml → base_config.mlp"),
        ("MLP dropout",         "0.3",       "stage2_config.yaml → base_config.mlp"),
        ("grad clip norm",      "1.0",       "anil.py / fomaml.py / maml.py"),
    ],
    col_widths=[2.2, 1.8, 3.0],
)

doc.save(OUTPUT_DESKTOP)
print("Saved:", OUTPUT_DESKTOP)
