from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/annita/Desktop/Thesis/Signal_Fusion_with_Meta-learners-/reports/notebook_pipeline_reference.docx"

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr_cells[i], "1F3864")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = val
            shade_cell(cells[ci], bg)
            cells[ci].paragraphs[0].paragraph_format.space_after = Pt(2)

    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()

# ── Title ──────────────────────────────────────────────────────────────────
add_heading("Notebook Pipeline Reference", level=1)
doc.add_paragraph(
    "This document describes the full pipeline of Jupyter notebooks in the "
    "Signal Fusion with Meta-Learners thesis project, including file inputs/outputs "
    "and inter-notebook dependencies."
)

# ── TABLE 1: File I/O ──────────────────────────────────────────────────────
add_heading("Table 1: Notebook File I/O", level=2)
add_heading("Data Preparation", level=3)

prep_headers = ["Notebook", "Loads", "Outputs"]
prep_rows = [
    ("01_prepare_contract_year",
     "Data/raw/dim_ci_stso_contracts_metadata.csv",
     "Data/raw/contract_year_final.csv"),
    ("02_prepare_esg",
     "/Users/annita/Desktop/Thesis/tables/moodys/ESG_Data_AXEF.csv",
     "tables/joining/esg_yearly.csv"),
    ("03_prepare_financials",
     "/Users/annita/Desktop/Thesis/tables/financials/moodys_financials.csv",
     "tables/joining/financials_clean.csv"),
    ("04_prepare_indexes",
     "Data/Index data/PPI.csv, LPI_Historical_Aligned.csv [⚠ hardcoded to Thomas's machine]",
     "Data/index data/macro_indexs.csv"),
    ("05_prepare_news",
     "/Users/annita/Desktop/Thesis/tables/suppliers/news_cleaned_llm.csv",
     "tables/joining/news.csv"),
    ("06_prepare_stocks",
     "Data/Stocks/Stock_Monthly_Trading.csv, Stock_Monthly_Market_Cap.csv, Stocks_beta_correl_etc.csv, Stocks_Monthly_Shares.csv, Stocks_Monthly_ClosingPrices.csv [⚠ hardcoded to Thomas's machine]",
     "Data/processed/stock_view.csv"),
    ("07_joining",
     "Data/raw/contract_year_final.csv, dim_supplier.csv, financials_clean.csv, esg_yearly.csv, news.csv, Stock_view.csv, macro_indexs.csv",
     "Data/processed/contract_final_no_spend.csv, gold_labels_manual.csv, contract_with_gold.csv, gold_labels_stage2_test.csv, contract_with_gold_stage2_test.csv"),
]
add_table(prep_headers, prep_rows, col_widths=[1.6, 2.8, 2.8])

add_heading("EDA", level=3)
eda_rows = [
    ("08_contracts_eda",
     "Data/raw/contract_year_final.csv",
     "Data/processed/contracts_eda_structure.csv, contracts_eda_missingness.csv, contracts_eda_imputation_candidates.csv"),
    ("09_esg_eda",
     "Data/raw/esg_yearly.csv",
     "Data/processed/esg_eda_structure.csv, esg_eda_missingness.csv, esg_eda_numeric_summary.csv, esg_eda_categorical_summary.csv, esg_eda_imputation_candidates.csv, esg_eda_feature_decisions.csv, esg_eda_high_corr_pairs.csv"),
    ("10_financials_eda",
     "Data/raw/financials_clean.csv",
     "Data/processed/eda_financial_screening.csv, eda_financial_snorkel_candidates.csv, eda_financial_imputation_decisions.csv, eda_financial_feature_decisions.csv, eda_financial_high_corr_pairs.csv"),
    ("11_indexes_eda",
     "Data/macro_indexs.csv [⚠ hardcoded to Thomas's machine]",
     "None found"),
    ("12_news_eda",
     "Data/raw/news.csv",
     "Data/processed/news_eda_structure.csv, news_eda_missingness.csv, news_eda_numeric_summary.csv, news_eda_categorical_summary.csv, news_eda_imputation_candidates.csv, news_eda_feature_decisions.csv, news_eda_high_corr_pairs.csv"),
    ("13_stocks_eda",
     "Data/raw/Stock_view.csv",
     "Data/processed/stocks_eda_structure.csv, stocks_eda_missingness.csv, stocks_eda_numeric_summary.csv, stocks_eda_categorical_summary.csv, stocks_eda_imputation_candidates.csv, stocks_eda_feature_decisions.csv, stocks_eda_high_corr_pairs.csv"),
]
add_table(prep_headers, eda_rows, col_widths=[1.6, 2.8, 2.8])

add_heading("Feature Engineering & Labeling", level=3)
label_rows = [
    ("14_feature_engineering",
     "Data/processed/contracts_with_features.csv",
     "Data/processed/contracts_with_features.csv (pruned), Data/processed/feature_pruning_audit.csv"),
    ("15_snorkel_labeling",
     "Data/processed/contracts_with_features.csv",
     "Data/processed/contract_with_features_labeled.csv, contract_snorkel_labeled.csv (compatibility copy), reports/tables/snorkel_lf_summary.csv, snorkel_group_summary.csv"),
    ("15.1_prepare_gold_labels",
     "Data/processed/contract_with_features_labeled.csv + gold labels hardcoded as Python dicts",
     "Data/processed/gold_labels_clean.csv, gold_labels_debug_augmented.csv"),
]
add_table(prep_headers, label_rows, col_widths=[1.6, 2.8, 2.8])

add_heading("Modeling", level=3)
model_rows = [
    ("16_stage1_baselines",
     "Data/processed/contract_with_features_labeled_with_gold.csv",
     "models/stage_1/{condition}/mlp_pretrained.pt, mlp_pretrained_preprocessor.joblib, mlp_pretrained_history.csv, tabular_results.csv, xgboost_feature_importance.csv; reports/tables/stage1_layer_a/b_vNNN.csv; reports/figures/"),
    ("17_stage2_meta_learning",
     "Data/processed/contract_with_features_labeled_with_gold.csv; models/stage_1/{init_name}/mlp_pretrained.pt + preprocessor.joblib; models/stage_2/experiments/experiment_summary.csv (if exists)",
     "models/stage_2/experiments/{id}/ — {method}_initialized.pt, metrics.csv, predictions.csv, history.csv, department_summary.csv, department_validity.csv, resolved_config.yaml, stage2_result_summary.json; experiment_summary.csv"),
    ("17.1_stage2_analysis_dashboard",
     "models/stage_2/experiments/experiment_summary.csv; per-experiment metrics.csv, predictions.csv, history.csv; optionally models/stage_1/stage1_experiment_summary.csv",
     "reports/stage2/table_D_stage2_main_benchmark.csv, table_E_adaptation_dynamics.csv, table_F_support_size_sensitivity.csv, table_G_init_ablation.csv, table_H_leaderboard.csv"),
]
add_table(prep_headers, model_rows, col_widths=[1.6, 2.8, 2.8])

add_heading("Utilities & Diagnostics", level=3)
util_rows = [
    ("18_dataset_diagnostic_report",
     "Data/processed/contract_with_features_labeled.csv (fallback: contracts_with_features.csv)",
     "reports/diagnostics/*.csv (dept distribution, gold distribution, coverage, validity, LF votes, missingness, feature sanity); reports/diagnostics/figures/*.png"),
    ("stage2_input_prep",
     "Data/processed/contract_with_features_labeled.csv; Data/processed/gold_labels_stage2_test.csv",
     "Exploratory only — no file saves"),
]
add_table(prep_headers, util_rows, col_widths=[1.6, 2.8, 2.8])

# ── TABLE 2: Dependencies (no EDA) ────────────────────────────────────────
add_heading("Table 2: Notebook Dependencies (EDA excluded)", level=2)

dep_headers = ["Notebook", "Created By", "Used By", "Contains"]
dep_rows = [
    ("01_prepare_contract_year",
     "dim_ci_stso_contracts_metadata.csv (external)",
     "07_joining",
     "Expands contract metadata into one row per contract-year; repairs missing department assignments"),
    ("02_prepare_esg",
     "Moody's ESG export (external)",
     "07_joining",
     "Cleans and structures ESG scores per supplier per year"),
    ("03_prepare_financials",
     "Moody's financials export (external)",
     "07_joining",
     "Cleans financial risk metrics (EBIT, solvency, profit margin) per supplier"),
    ("04_prepare_indexes",
     "PPI, LPI index files (external)",
     "07_joining",
     "Aligns macro-economic indices to observation years"),
    ("05_prepare_news",
     "LLM-classified news CSV (external)",
     "07_joining",
     "Structures news sentiment and frequency features per supplier"),
    ("06_prepare_stocks",
     "Stock files (external)",
     "07_joining",
     "Computes stock volatility, market cap, beta, price trend features per supplier"),
    ("07_joining",
     "01, 02, 03, 04, 05, 06",
     "14_feature_engineering, 15_snorkel_labeling, stage2_input_prep",
     "Merges all domain tables into the master feature dataset"),
    ("14_feature_engineering",
     "07_joining",
     "15_snorkel_labeling, 16_stage1_baselines, 17_stage2_meta_learning",
     "Applies domain pruning rules to drop leaky/redundant features; saves pruning audit log"),
    ("15_snorkel_labeling",
     "14_feature_engineering",
     "15.1_prepare_gold_labels, stage2_input_prep, 16_stage1_baselines, 18_dataset_diagnostic_report",
     "Applies ~40 Snorkel labeling functions; generates soft label renegotiation_prob for all contracts"),
    ("15.1_prepare_gold_labels",
     "15_snorkel_labeling (features), hardcoded dicts (gold labels)",
     "stage2_input_prep, 16_stage1_baselines",
     "Inlines manually curated gold labels per department; merges with feature table"),
    ("stage2_input_prep",
     "15_snorkel_labeling, 15.1_prepare_gold_labels",
     "16_stage1_baselines, 17_stage2_meta_learning",
     "Merges weak-labeled features with gold labels into the unified dataset used by both stages"),
    ("16_stage1_baselines",
     "stage2_input_prep",
     "17_stage2_meta_learning (MLP weights), 17.1 (reference metrics)",
     "Trains ElasticNet, XGBoost, Mean Predictor, MLP under 3 conditions (weak-only, gold-only, hybrid); evaluates on weak and gold labels"),
    ("17_stage2_meta_learning",
     "stage2_input_prep (data), 16_stage1_baselines (MLP + preprocessor)",
     "17.1_stage2_analysis_dashboard",
     "Builds department-level few-shot episodes; meta-trains ANIL, FOMAML, MAML; evaluates on held-out Logistics"),
    ("17.1_stage2_analysis_dashboard",
     "17_stage2_meta_learning",
     "Human / thesis writing",
     "Aggregates all Stage 2 results; produces benchmark, sensitivity, and leaderboard tables"),
    ("18_dataset_diagnostic_report",
     "15_snorkel_labeling",
     "Human / thesis writing",
     "Full dataset quality audit: department distribution, gold label coverage, class balance, Snorkel diagnostics, feature missingness"),
]
add_table(dep_headers, dep_rows, col_widths=[1.6, 1.5, 1.9, 2.2])

# ── Notes ──────────────────────────────────────────────────────────────────
add_heading("Notes", level=2)
notes = [
    "Notebooks 04 and 06 have paths hardcoded to Thomas's machine (/Users/Thomas/...) and will fail on another machine without updating the paths.",
    "15.1_prepare_gold_labels has no raw file input for gold labels — they are hardcoded as Python dicts directly in the notebook.",
    "contract_with_features_labeled_with_gold.csv (used by 16 and 17) is produced by stage2_input_prep merging the outputs of 15 and 15.1.",
    "EDA notebooks (08–13) are dead ends — their outputs are for human review only, not consumed by any downstream notebook.",
    "stage2_input_prep is the critical bridge that produces the single CSV consumed by both Stage 1 and Stage 2.",
    "16_stage1_baselines is the only notebook that produces model artifacts used by another notebook (the pretrained MLP loaded by 17).",
]
for note in notes:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(note)

doc.save(OUTPUT)
print("Saved:", OUTPUT)
